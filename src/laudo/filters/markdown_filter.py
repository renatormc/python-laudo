from __future__ import annotations

from docx import Document
from docx.shared import Pt
import re
from typing import TYPE_CHECKING
if TYPE_CHECKING:
    from docxtpl import Subdoc
    from laudo.core import RenderEnv
from uuid import uuid4


class MarkdownFilter:
    def __init__(self, renv: RenderEnv):
        self.jenv = renv

    def __call__(self, md_text: str) -> Subdoc:
        tempdocx = self.jenv.temp_folder / f"{uuid4().hex}.docx"

  
        document = Document()

        # Pre-compile simple patterns
        bold_re = re.compile(r"\*\*(.*?)\*\*")
        italic_re = re.compile(r"\*(.*?)\*")

        lines = md_text.split("\n")
        list_mode = False
        current_list = None

        for line in lines:
            stripped = line.strip()

            # ---------- HEADINGS ----------
            if stripped.startswith("### "):
                heading = document.add_heading(stripped[4:], level=3)
                heading.paragraph_format.space_before = Pt(0)
                heading.paragraph_format.space_after = Pt(0)
                list_mode = False
                continue
            elif stripped.startswith("## "):
                heading = document.add_heading(stripped[3:], level=2)
                heading.paragraph_format.space_before = Pt(0)
                heading.paragraph_format.space_after = Pt(0)
                list_mode = False
                continue
            elif stripped.startswith("# "):
                heading = document.add_heading(stripped[2:], level=1)
                heading.paragraph_format.space_before = Pt(0)
                heading.paragraph_format.space_after = Pt(0)
                list_mode = False
                continue

            # ---------- LIST ITEM ----------
            if stripped.startswith("- "):
                text = stripped[2:]

                if not list_mode:
                    # start new list
                    current_list = document.add_paragraph(style="Normal")
                    list_mode = True
                else:
                    # continue list
                    current_list = document.add_paragraph(style="Normal")

                # Add bullet symbol
                current_list.add_run("\u2022 ")
                
                # Parse formatting in list item text
                tokens = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
                for token in tokens:
                    if bold_re.match(token):
                        current_list.add_run(bold_re.findall(token)[0]).bold = True
                    elif italic_re.match(token):
                        current_list.add_run(italic_re.findall(token)[0]).italic = True
                    else:
                        current_list.add_run(token)

                # Remove extra spacing for list items
                current_list.paragraph_format.space_before = Pt(0)
                current_list.paragraph_format.space_after = Pt(0)
                # Keep document `Normal` style but make it look like a bulleted list
                current_list.paragraph_format.left_indent = Pt(18)
                current_list.paragraph_format.first_line_indent = Pt(-9)
                continue

            # ---------- NORMAL PARAGRAPH ----------
            list_mode = False
            if stripped == "":
                continue  # skip blank lines instead of adding empty paragraphs

            # Apply bold and italic manually
            paragraph = document.add_paragraph(style="Normal")
            
            # Remove extra spacing before and after paragraph
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            
            idx = 0
            temp = line

            # Tokenize by bold/italic delimiters
            tokens = re.split(r"(\*\*.*?\*\*|\*.*?\*)", temp)

            for token in tokens:
                if bold_re.match(token):
                    paragraph.add_run(bold_re.findall(token)[0]).bold = True
                elif italic_re.match(token):
                    paragraph.add_run(italic_re.findall(token)[0]).italic = True
                else:
                    paragraph.add_run(token)

        document.save(str(tempdocx))
        return self.jenv.tpl.new_subdoc(str(tempdocx))
